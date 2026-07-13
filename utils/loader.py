from __future__ import annotations

import re
from pathlib import Path
from typing import List

from config import KNOWLEDGE_DIR


class DocumentLoader:
    """Loads and normalizes supported documents from the knowledge base.

    Heavy parsing libraries (bs4, pypdf, docx) are imported lazily so that
    cold-start time is not penalised by libraries whose file types may not
    exist in the current knowledge base.
    """

    def __init__(self, base_dir: Path | None = None) -> None:
        self.base_dir = base_dir or KNOWLEDGE_DIR
        self.supported_extensions = {".pdf", ".docx", ".md", ".markdown", ".html", ".htm", ".txt"}

    def load_documents(self) -> List[dict]:
        documents: List[dict] = []
        if not self.base_dir.exists():
            return documents

        for path in sorted(self.base_dir.rglob("*")):
            if not path.is_file():
                continue
            if path.suffix.lower() not in self.supported_extensions:
                continue

            try:
                documents.extend(self._load_single_document(path))
            except Exception as exc:
                import logging
                logging.warning(f"Skipped corrupted file: {path} ({exc})")

        return documents

    def _load_single_document(self, path: Path) -> List[dict]:
        extension = path.suffix.lower()
        if extension == ".pdf":
            return self._load_pdf(path)
        if extension == ".docx":
            return [self._build_document_record(path, self._extract_docx_text(path), "docx")]
        if extension in {".md", ".markdown"}:
            text = path.read_text(encoding="utf-8")
            return [self._build_document_record(path, self._clean_markdown_text(text), "markdown")]
        if extension in {".html", ".htm"}:
            return [self._build_document_record(path, self._extract_html_text(path), "html")]
        if extension == ".txt":
            text = path.read_text(encoding="utf-8")
            return [self._build_document_record(path, self._clean_text(text), "txt")]
        return []

    def _load_pdf(self, path: Path) -> List[dict]:
        # Lazy import: only pay the pypdf initialisation cost when a PDF is found.
        from pypdf import PdfReader  # noqa: PLC0415

        reader = PdfReader(str(path))
        pages: List[dict] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            cleaned = self._clean_text(text)
            if cleaned.strip():
                pages.append(self._build_document_record(path, cleaned, "pdf", page_number=page_number))

        if not pages:
            pages.append(self._build_document_record(path, "", "pdf"))
        return pages

    def _extract_docx_text(self, path: Path) -> str:
        # Lazy import: only pay the python-docx initialisation cost when a DOCX is found.
        from docx import Document as DocxDocument  # noqa: PLC0415

        document = DocxDocument(str(path))
        parts: List[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
        return self._clean_text("\n".join(parts))

    def _extract_html_text(self, path: Path) -> str:
        # Lazy import: only pay the BeautifulSoup initialisation cost when an HTML file is found.
        from bs4 import BeautifulSoup  # noqa: PLC0415

        html = path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text("\n", strip=True)
        return self._clean_text(text)

    def _build_document_record(
        self,
        path: Path,
        content: str,
        document_type: str,
        page_number: int | None = None,
    ) -> dict:
        builder_name, project_name = self._detect_metadata(content)
        return {
            "source": str(path.relative_to(self.base_dir)).replace("\\", "/"),
            "content": content.strip(),
            "document_name": path.name,
            "document_path": str(path.resolve()),
            "file_type": document_type,
            "page_number": page_number,
        }

    def _detect_metadata(self, content: str) -> tuple[str | None, str | None]:
        return None, None

    def _clean_markdown_text(self, text: str) -> str:
        cleaned = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
        cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)
        cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
        cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
        cleaned = re.sub(r"^\s*[-*+]\s*", "", cleaned, flags=re.MULTILINE)
        cleaned = re.sub(r"\[(.*?)\]\([^)]*\)", r"\1", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return self._clean_text(cleaned)

    def _clean_text(self, text: str) -> str:
        normalized = re.sub(r"\s+", " ", text.replace("\r\n", "\n").replace("\r", "\n"))
        return normalized.strip()
